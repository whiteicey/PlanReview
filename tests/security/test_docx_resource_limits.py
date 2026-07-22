from __future__ import annotations

from dataclasses import replace
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile

from docx import Document
import pytest

from app.domain.exceptions import DocxResourceLimitError, UnsafeDocxPackageError
from app.parsers.docx_parser import DocxParser
from app.settings import Settings
from app.storage.case_files import _safe_archive_member_name, validate_docx_package


CORE = {
    "[Content_Types].xml": b"<Types/>",
    "_rels/.rels": b"<Relationships/>",
    "word/document.xml": b"<w:document xmlns:w='urn:w'><w:body/></w:document>",
}


def _archive(path: Path, extra=None, *, compression=ZIP_STORED) -> Path:
    with ZipFile(path, "w", compression=compression) as archive:
        for name, value in {**CORE, **(extra or {})}.items():
            archive.writestr(name, value)
    return path


@pytest.mark.parametrize("name", ["../escape.xml", "/absolute.xml", "C:/drive.xml"])
def test_unsafe_archive_paths_are_rejected_without_echo(tmp_path, name):
    path = _archive(tmp_path / "bad.docx", {name: b"x"})
    with pytest.raises(UnsafeDocxPackageError) as error:
        validate_docx_package(path)
    assert name not in str(error.value)


def test_backslash_archive_member_is_rejected(tmp_path):
    with pytest.raises(UnsafeDocxPackageError):
        _safe_archive_member_name("word\\bad.xml")


def test_safe_unknown_office_extension_directory_is_allowed(tmp_path):
    path = _archive(tmp_path / "extension.docx", {"vendorExtension/data.xml": b"safe"})
    validate_docx_package(path)


@pytest.mark.parametrize("name", ["word/embeddings/object.bin", "word/activeX/control.bin", "word/vbaProject.bin"])
def test_dangerous_embedded_objects_are_rejected(tmp_path, name):
    with pytest.raises(UnsafeDocxPackageError):
        validate_docx_package(_archive(tmp_path / "embedded.docx", {name: b"x"}))


def test_member_count_total_size_member_size_and_ratio_limits(tmp_path):
    many = _archive(tmp_path / "many.docx", {"custom/a": b"a"})
    with pytest.raises(DocxResourceLimitError):
        validate_docx_package(many, replace(Settings(), max_zip_members=3))

    total = _archive(tmp_path / "total.docx", {"custom/a": b"12345"})
    actual_total = sum(info.file_size for info in ZipFile(total).infolist())
    validate_docx_package(total, replace(Settings(), max_zip_uncompressed_bytes=actual_total))
    with pytest.raises(DocxResourceLimitError):
        validate_docx_package(total, replace(Settings(), max_zip_uncompressed_bytes=actual_total - 1))

    member = _archive(tmp_path / "member.docx", {"custom/large": b"x" * 100})
    with pytest.raises(DocxResourceLimitError):
        validate_docx_package(member, replace(Settings(), max_zip_member_bytes=99))

    ratio = _archive(
        tmp_path / "ratio.docx", {"custom/compressed": b"0" * 20_000}, compression=ZIP_DEFLATED
    )
    with pytest.raises(DocxResourceLimitError):
        validate_docx_package(ratio, replace(Settings(), max_zip_compression_ratio=10))


def test_parser_character_and_cell_boundaries(tmp_path):
    character_path = tmp_path / "characters.docx"
    document = Document()
    document.add_paragraph("12345")
    document.save(character_path)
    DocxParser(replace(Settings(), max_document_characters=5)).parse(character_path)
    with pytest.raises(DocxResourceLimitError):
        DocxParser(replace(Settings(), max_document_characters=4)).parse(character_path)

    cell_path = tmp_path / "cells.docx"
    document = Document()
    document.add_table(rows=1, cols=2)
    document.save(cell_path)
    DocxParser(replace(Settings(), max_table_cells=2)).parse(cell_path)
    with pytest.raises(DocxResourceLimitError):
        DocxParser(replace(Settings(), max_table_cells=1)).parse(cell_path)
