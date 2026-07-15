"""The web review path must run the rule engine, and degrade honestly."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient


def _capacity_docx() -> bytes:
    document = Document()
    document.add_heading("产能与生产预测", level=1)
    table = document.add_table(rows=3, cols=7)
    header = ["参数名称", "数值", "单位", "对象", "时间/阶段", "统计口径", "条件"]
    peak = ["高峰产量", "230", "万m³/d", "气田_A", "达产期", "设计工况", ""]
    capacity = ["地面处理能力", "200", "万m³/d", "气田_A", "设计期", "设计能力", ""]
    for column, text in enumerate(header):
        table.rows[0].cells[column].text = text
    for column, text in enumerate(peak):
        table.rows[1].cells[column].text = text
    for column, text in enumerate(capacity):
        table.rows[2].cells[column].text = text
    document.add_paragraph("高峰产量超过处理能力，请复核。")
    payload = BytesIO()
    document.save(payload)
    return payload.getvalue()


def _client(monkeypatch, tmp_path: Path) -> TestClient:
    monkeypatch.setenv("REVIEW_STORAGE_ROOT", str(tmp_path / "storage"))
    from app.settings import get_settings

    get_settings.cache_clear()
    from app.main import app

    return TestClient(app)


def _upload(client: TestClient) -> str:
    response = client.post(
        "/api/cases",
        files={"file": ("方案.docx", _capacity_docx(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 201
    return response.json()["case_id"]


def _load_real_ruleset():
    # Reuse the actual external ruleset if configured; skip when absent so the
    # test never fakes a pass.
    import pytest

    from app.rules.ruleset import RulesetNotConfigured, load_active_ruleset

    try:
        return load_active_ruleset()
    except RulesetNotConfigured as exc:
        pytest.skip(f"外部规则库未配置: {exc}")


def test_review_runs_rule_engine_when_ruleset_loaded(monkeypatch, tmp_path):
    loaded = _load_real_ruleset()
    from app.api import routes

    monkeypatch.setattr(routes, "_active_ruleset", lambda: loaded)
    client = _client(monkeypatch, tmp_path)
    case_id = _upload(client)

    review = client.post(f"/api/cases/{case_id}/review")
    assert review.status_code == 201
    body = review.json()
    assert body["rules_loaded"] is True
    assert body["rule_count"] == len(loaded.rules)

    findings = client.get(f"/api/cases/{case_id}/findings").json()
    # At least one finding must come from a real rule (not only the Mock LLM).
    rule_findings = [f for f in findings if f["rule_id"]]
    assert rule_findings, findings
    assert any(f["rule_id"] == "CAPACITY-001" for f in rule_findings)


def test_review_degrades_when_ruleset_not_configured(monkeypatch, tmp_path):
    from app.api import routes

    monkeypatch.setattr(routes, "_active_ruleset", lambda: None)
    client = _client(monkeypatch, tmp_path)
    case_id = _upload(client)

    review = client.post(f"/api/cases/{case_id}/review")
    assert review.status_code == 201
    body = review.json()
    assert body["rules_loaded"] is False
    assert body["rule_count"] == 0

    findings = client.get(f"/api/cases/{case_id}/findings").json()
    # Degraded path still runs the Mock LLM, but no rule-derived findings.
    assert all(f["rule_id"] is None for f in findings)


def test_ruleset_status_and_reload_endpoints(monkeypatch, tmp_path):
    loaded = _load_real_ruleset()
    from app.api import routes

    # Start from an empty cache so status reflects reload behaviour, not import order.
    routes._reset_ruleset_cache()
    monkeypatch.setattr(routes, "load_active_ruleset", lambda root=None: loaded)
    client = _client(monkeypatch, tmp_path)

    reload_response = client.post("/api/ruleset/reload", json={})
    assert reload_response.status_code == 200
    body = reload_response.json()
    assert body["loaded"] is True
    assert body["rule_count"] == len(loaded.rules)
    assert body["root"]

    status = client.get("/api/ruleset").json()
    assert status["loaded"] is True
    assert status["rule_count"] == len(loaded.rules)


def test_ruleset_reload_reports_unconfigured_without_500(monkeypatch, tmp_path):
    from app.api import routes
    from app.rules.ruleset import RulesetNotConfigured

    routes._reset_ruleset_cache()

    def _raise(root=None):
        raise RulesetNotConfigured("找不到示例数据包")

    monkeypatch.setattr(routes, "load_active_ruleset", _raise)
    client = _client(monkeypatch, tmp_path)

    reload_response = client.post("/api/ruleset/reload", json={})
    assert reload_response.status_code == 200
    body = reload_response.json()
    assert body["loaded"] is False
    assert body["rule_count"] == 0
    assert body["root"] is None
