"""End-to-end acceptance path over a self-contained, synthetic document.

Exercises the whole kernel without the external DEMO package: parse a DOCX ->
extract facts with SourceSpans -> run the three-valued rule engine -> Mock LLM
review under the evidence gate -> reconcile -> Findings that trace back to real
evidence. Also asserts the source tree contains no eval/exec.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document

from app.domain.enums import OnMissing, RuleStatus, Severity
from app.domain.schemas import RuleDefinition
from app.llm.mock import MockProvider
from app.parsers.docx_parser import DocxParser
from app.review.pipeline import ReviewPipeline

ROOT = Path(__file__).resolve().parents[2]


def _capacity_docx(tmp_path: Path) -> Path:
    document = Document()
    document.add_heading("产能与生产预测", level=1)
    table = document.add_table(rows=3, cols=7)
    header = ["参数名称", "数值", "单位", "对象", "时间/阶段", "统计口径", "条件"]
    for column, text in enumerate(header):
        table.rows[0].cells[column].text = text
    peak = ["高峰产量", "230", "万m³/d", "气田_A", "达产期", "设计工况", ""]
    capacity = ["地面处理能力", "200", "万m³/d", "气田_A", "设计期", "设计能力", ""]
    for column, text in enumerate(peak):
        table.rows[1].cells[column].text = text
    for column, text in enumerate(capacity):
        table.rows[2].cells[column].text = text
    document.add_paragraph("高峰产量超过处理能力，请复核。")
    path = tmp_path / "capacity.docx"
    document.save(path)
    return path


def _capacity_rule() -> RuleDefinition:
    return RuleDefinition(
        rule_id="CAPACITY-001",
        version="acceptance",
        name="高峰产量不超过地面处理能力",
        category="cross_domain",
        severity=Severity.HIGH,
        operator="less_or_equal",
        on_missing=OnMissing.UNKNOWN,
        params={"left": "高峰产量", "right": "地面处理能力", "parameter": "高峰产量"},
    )


def test_end_to_end_acceptance_path(tmp_path):
    path = _capacity_docx(tmp_path)
    document = DocxParser().parse(path, document_id="ACCEPT")

    run = ReviewPipeline().run("ACCEPT", [document], [_capacity_rule()], MockProvider())

    # Facts carry SourceSpans.
    assert run.facts
    assert all(fact.source_span_id for fact in run.facts)
    span_ids = {span_id for span_id in run.evidence_text_hashes}

    # The rule fires on the real values (230 万 > 200 万 after normalization).
    capacity_results = [r for r in run.rule_results if r.rule_id == "CAPACITY-001"]
    assert capacity_results
    assert capacity_results[0].status is RuleStatus.FAIL

    # Findings exist, each with evidence tracing to real spans.
    assert run.findings
    for finding in run.findings:
        assert finding.evidence_span_ids
        assert set(finding.evidence_span_ids).issubset(span_ids)


def test_source_tree_contains_no_eval_or_exec():
    offenders = []
    for base in (ROOT / "app", ROOT / "scripts"):
        for py in base.rglob("*.py"):
            if "__pycache__" in py.parts:
                continue
            text = py.read_text(encoding="utf-8")
            if "eval(" in text or "exec(" in text:
                offenders.append(str(py.relative_to(ROOT)))
    assert not offenders, offenders


def test_acceptance_command_is_documented():
    readme = (ROOT / "README.md").read_text(encoding="utf-8")
    assert "python -m pytest -q" in readme
    assert "AI 初审结果" in readme
