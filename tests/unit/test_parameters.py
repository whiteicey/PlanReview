from pathlib import Path

from docx import Document

from app.domain.enums import ExtractionMethod
from app.parsers.docx_parser import DocxParser


def make_fact_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("附件A关键参数表", level=1)
    table = doc.add_table(rows=3, cols=7)
    for index, header in enumerate(
        ["参数名称", "数值", "单位", "对象", "时间/阶段", "统计口径", "条件"]
    ):
        table.cell(0, index).text = header
    for index, value in enumerate(
        ["开发井总数", "36", "口", "全区", "全生命周期", "累计", "基准方案"]
    ):
        table.cell(1, index).text = value
    for index, value in enumerate(
        ["高峰产量", "220", "万m³/d", "全区", "达产期", "日峰值", "稳产条件"]
    ):
        table.cell(2, index).text = value
    doc.add_paragraph("建设周期为30个月，投产时间：2028年03月。")
    doc.save(path)


def test_extracts_full_table_key(tmp_path: Path) -> None:
    from app.extraction.parameters import extract_parameter_facts

    path = tmp_path / "facts.docx"
    make_fact_docx(path)
    facts = extract_parameter_facts(DocxParser().parse(path, "D1"), "V1")

    total = next(fact for fact in facts if fact.raw_name == "开发井总数")
    assert (total.raw_value, total.normalized_value) == ("36", 36.0)
    assert (total.raw_unit, total.subject, total.time_scope, total.statistical_scope) == (
        "口",
        "全区",
        "全生命周期",
        "累计",
    )
    assert total.condition == "基准方案"
    assert total.source_version == "V1"
    assert total.extraction_method is ExtractionMethod.TABLE


def test_extracts_body_number_and_unit(tmp_path: Path) -> None:
    from app.extraction.parameters import extract_parameter_facts

    path = tmp_path / "facts.docx"
    make_fact_docx(path)
    parsed = DocxParser().parse(path, "D1")
    facts = extract_parameter_facts(parsed)

    cycle = next(fact for fact in facts if fact.raw_name == "建设周期")
    assert cycle.normalized_value == 30.0
    assert cycle.raw_unit == "个月"
    assert cycle.source_span_id == next(
        span.span_id for span in parsed.paragraphs if "建设周期" in span.text
    )
    assert cycle.extraction_method is ExtractionMethod.REGEX


def test_retains_conflicting_table_and_body_occurrences(tmp_path: Path) -> None:
    from app.extraction.parameters import extract_parameter_facts

    path = tmp_path / "conflicting-facts.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "参数名称"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "开发井总数"
    table.cell(1, 1).text = "36"
    doc.add_paragraph("开发井总数为40口。")
    doc.save(path)

    parsed = DocxParser().parse(path, "D2")
    facts = extract_parameter_facts(parsed)
    wells = [fact for fact in facts if fact.raw_name == "开发井总数"]

    assert [(fact.raw_value, fact.source_span_id, fact.extraction_method) for fact in wells] == [
        ("36", "D2:t:0:1:1", ExtractionMethod.TABLE),
        ("40", "D2:p:0", ExtractionMethod.REGEX),
    ]
    assert len({fact.fact_id for fact in wells}) == 2


def test_does_not_fill_missing_dimensions(tmp_path: Path) -> None:
    from app.extraction.parameters import extract_parameter_facts

    path = tmp_path / "sparse.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    for index, header in enumerate(["参数名称", "数值", "单位"]):
        table.cell(0, index).text = header
    for index, value in enumerate(["设计能力", "1,250", "万吨/年"]):
        table.cell(1, index).text = value
    doc.save(path)

    [fact] = extract_parameter_facts(DocxParser().parse(path, "D3"))

    assert fact.normalized_value == 1250.0
    assert fact.raw_unit == "万吨/年"
    assert fact.subject is None
    assert fact.time_scope is None
    assert fact.statistical_scope is None
    assert fact.condition is None


def test_flattens_multi_row_merged_headers_and_preserves_separate_time_stage(
    tmp_path: Path,
) -> None:
    from app.extraction.parameters import extract_parameter_facts

    path = tmp_path / "merged-headers.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=4)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "参数"
    table.cell(0, 2).merge(table.cell(0, 3)).text = "时间维度"
    for index, header in enumerate(["名称", "数值", "时间", "阶段"]):
        table.cell(1, index).text = header
    for index, value in enumerate(["高峰产量", "220", "2028年", "达产期"]):
        table.cell(2, index).text = value
    doc.save(path)

    facts = extract_parameter_facts(DocxParser().parse(path, "D4"))

    [fact] = [fact for fact in facts if fact.raw_name == "高峰产量"]
    assert (fact.raw_value, fact.normalized_value) == ("220", 220.0)
    assert fact.time_scope == "时间=2028年;阶段=达产期"
    assert fact.source_span_id == "D4:t:0:2:1"


def test_does_not_extract_partial_date_like_body_value(tmp_path: Path) -> None:
    from app.extraction.parameters import extract_parameter_facts

    path = tmp_path / "date-like-body.docx"
    doc = Document()
    doc.add_paragraph("投产时间：2028年03月，建设周期为30个月。")
    doc.save(path)

    facts = extract_parameter_facts(DocxParser().parse(path, "D5"))

    assert not any(fact.raw_name == "投产时间" for fact in facts)
    cycle = next(fact for fact in facts if fact.raw_name == "建设周期")
    assert (cycle.raw_value, cycle.raw_unit) == ("30", "个月")
