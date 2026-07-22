from __future__ import annotations

from copy import deepcopy
from hashlib import sha256
from io import BytesIO
import json
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook
import pytest
from pydantic import TypeAdapter
from sqlalchemy import func, select
from sqlalchemy.exc import IntegrityError

from app.domain.schemas import Finding
from app.main import app
from app.persistence.db import DatabaseRuntime
from app.persistence.models import CaseRecord, FindingORM, ReviewRunORM, RuleResultORM
from app.persistence.repository import ReviewRepository
from app.review.pipeline import ReviewRun
from app.settings import get_settings
from app.storage.case_files import StoredFile


FIXTURE_PATH = Path(__file__).resolve().parents[1] / "fixtures" / "v12_real_terminal_payload.json"
RUN_ADAPTER = TypeAdapter(ReviewRun)
EXPECTED_COUNTS = {
    "distinct_rule_id": 18,
    "rule_results": 29,
    "rule_findings": 24,
    "checkpoint_ai_findings": 38,
    "final_findings": 25,
    "deduplication_records": 37,
    "batch_metrics": 5,
    "packet_ledger_entries": 504,
    "ai_candidate_ledger_entries": 76,
    "stage_records": 8,
    "facts": 16,
    "evidence_spans": 1719,
}


def _fixture() -> dict:
    return json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))


def _review_run(value: dict | None = None) -> ReviewRun:
    payload = _fixture()["review_run"] if value is None else value
    return RUN_ADAPTER.validate_python(payload)


def _business_hash(run: ReviewRun) -> str:
    payload = RUN_ADAPTER.dump_python(run, mode="json")
    encoded = json.dumps(
        payload,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")
    return sha256(encoded).hexdigest()


def _write_case_document(storage_root: Path, case_id: str) -> StoredFile:
    relative = Path("cases") / case_id / "documents" / "fixture.docx"
    target = storage_root / relative
    target.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Fixture", level=1)
    document.add_paragraph("Sanitized terminal persistence replay document.")
    document.save(target)
    digest = sha256(target.read_bytes()).hexdigest()
    return StoredFile(
        storage_relative_path=relative.as_posix(),
        sha256=digest,
        size=target.stat().st_size,
        safe_name="fixture.docx",
    )


def _create_running_case(runtime: DatabaseRuntime, run: ReviewRun) -> str:
    stored = _write_case_document(runtime.path.parent, run.case_id)
    session = runtime.session()
    repository = ReviewRepository(session)
    repository.save_case(CaseRecord(case_id=run.case_id, files=[stored]))
    repository.create_running_run(run.case_id, run.run_id)
    token = "22222222-2222-4222-8222-222222222222"
    assert repository.claim_running_run(run.run_id, token)
    session.close()
    return token


def _database_counts(runtime: DatabaseRuntime, run_id: str) -> tuple[int, int]:
    session = runtime.session()
    internal_id = session.scalar(
        select(ReviewRunORM.id).where(ReviewRunORM.run_id == run_id)
    )
    assert internal_id is not None
    finding_count = session.scalar(
        select(func.count()).select_from(FindingORM).where(
            FindingORM.review_run_id == internal_id
        )
    )
    result_count = session.scalar(
        select(func.count()).select_from(RuleResultORM).where(
            RuleResultORM.review_run_id == internal_id
        )
    )
    session.close()
    return int(finding_count or 0), int(result_count or 0)


def _xlsx_counts(content: bytes) -> tuple[int, int, list[str]]:
    workbook = load_workbook(BytesIO(content), data_only=False)
    findings = workbook["Findings"]
    rules = workbook["Rules"]
    finding_headers = [cell.value for cell in findings[2]]
    finding_id_column = finding_headers.index("finding_id") + 1
    finding_ids = [
        findings.cell(row=row, column=finding_id_column).value
        for row in range(3, findings.max_row + 1)
    ]
    return findings.max_row - 2, rules.max_row - 1, finding_ids


def test_v12_terminal_payload_fixture_invariants() -> None:
    fixture = _fixture()
    assert fixture["fixture_nature"] == (
        "基于真实Run现场、冻结DOCX、真实checkpoint数据和当前确定性规则重建的生产等价终态Payload"
    )
    assert fixture["not_a_direct_database_terminal_snapshot"] is True
    assert fixture["expected_counts"] == EXPECTED_COUNTS
    assert fixture["checkpoint_sizes"] == [8, 16, 24, 32, 38]
    assert len(fixture["checkpoint_findings"]) == 38
    assert len(fixture["rule_findings"]) == 24

    run = _review_run(fixture["review_run"])
    assert len(run.rule_results) == 29
    assert len({item.rule_id for item in run.rule_results}) == 18
    assert len(run.findings) == 25
    assert len(run.deduplication_records) == 37
    assert len(run.batch_metrics) == 5
    assert len(run.stage_records) == 8
    assert len(run.facts) == 16
    assert len(run.evidence_text_hashes) == 1719
    assert run.packet_lifecycle_ledger["ledger_entry_count"] == 504
    assert run.ai_candidate_lifecycle_ledger["ledger_entry_count"] == 76
    assert sum(item.parameter == "" for item in run.rule_results) == 2
    assert sum(item.parameter == "" for item in run.findings) == 2

    signatures = [
        json.dumps(item.model_dump(mode="json"), ensure_ascii=False, sort_keys=True)
        for item in run.rule_results
    ]
    assert len(signatures) == len(set(signatures))
    for rule_id in {item.rule_id for item in run.rule_results}:
        rows = [item for item in run.rule_results if item.rule_id == rule_id]
        if len(rows) > 1:
            scopes = {
                (item.parameter, tuple(item.evidence_span_ids))
                for item in rows
            }
            assert len(scopes) == len(rows)

    forbidden = json.dumps(fixture, ensure_ascii=False).casefold()
    assert "api_key" not in forbidden
    assert "authorization" not in forbidden
    assert "provider raw response" not in forbidden
    assert "d:\\" not in forbidden


def test_v12_real_terminal_payload_type_adapter_roundtrip_hash() -> None:
    first = _review_run()
    first_hash = _business_hash(first)
    dumped = RUN_ADAPTER.dump_python(first, mode="json")
    second = RUN_ADAPTER.validate_python(dumped)
    assert _business_hash(second) == first_hash


def test_v12_real_terminal_payload_replay_end_to_end(monkeypatch, tmp_path) -> None:
    storage = tmp_path / "storage"
    monkeypatch.setenv("REVIEW_STORAGE_ROOT", str(storage))
    get_settings.cache_clear()
    run = _review_run()
    fixture = _fixture()

    runtime = DatabaseRuntime(storage / "review.db")
    runtime.initialize()
    token = _create_running_case(runtime, run)
    session = runtime.session()
    repository = ReviewRepository(session)
    checkpoints = [Finding.model_validate(item) for item in fixture["checkpoint_findings"]]
    for size in fixture["checkpoint_sizes"]:
        repository.checkpoint_running_run(run, token, checkpoints[:size])
    repository.finish_running_run(run, token)
    session.close()
    runtime.dispose()

    restarted = DatabaseRuntime(storage / "review.db")
    restarted.initialize()
    session = restarted.session()
    repository = ReviewRepository(session)
    persisted = repository.get_run(run.run_id)
    assert persisted is not None
    assert persisted.final_status == "READY_FOR_HUMAN_REVIEW"
    assert len(persisted.findings) == 25
    assert len(persisted.rule_results) == 29
    assert len({item.rule_id for item in persisted.rule_results}) == 18
    assert persisted.packet_lifecycle_ledger["ledger_entry_count"] == 504
    assert persisted.ai_candidate_lifecycle_ledger["ledger_entry_count"] == 76
    assert len(persisted.rule_metrics) == 6
    assert len(persisted.batch_metrics) == 5
    assert len(persisted.stage_records) == 8
    assert len({item.finding_id for item in persisted.findings}) == 25
    assert _database_counts(restarted, run.run_id) == (25, 29)

    repository.finish_running_run(run, token)
    assert _database_counts(restarted, run.run_id) == (25, 29)
    changed = _review_run()
    changed.findings[0] = changed.findings[0].model_copy(update={"title": "Different payload"})
    with pytest.raises(ValueError, match="not owned"):
        repository.finish_running_run(changed, token)
    assert _database_counts(restarted, run.run_id) == (25, 29)
    session.close()
    restarted.dispose()

    with TestClient(app) as client:
        page = client.get("/")
        assert page.status_code == 200
        script = client.get("/app.js")
        assert script.status_code == 200
        assert "summary.finding_count" in script.text
        assert "findings.map(findingCard)" in script.text

        summary = client.get(f"/api/cases/{run.case_id}/runs/{run.run_id}")
        assert summary.status_code == 200
        assert summary.json()["finding_count"] == 25
        findings = client.get(f"/api/cases/{run.case_id}/runs/{run.run_id}/findings")
        assert findings.status_code == 200
        assert len(findings.json()) == 25
        diagnostics = client.get(
            f"/api/cases/{run.case_id}/runs/{run.run_id}/diagnostics"
        )
        assert diagnostics.status_code == 200
        integrity = diagnostics.json()["integrity"]
        assert integrity["finding_count"] == 25
        assert integrity["rule_result_count"] == 29
        assert integrity["distinct_rule_id_count"] == 18
        assert integrity["batch_count"] == 5

        exported = client.get(f"/api/cases/{run.case_id}/exports/xlsx")
        assert exported.status_code == 200
        assert _xlsx_counts(exported.content) == (
            25,
            29,
            [item.finding_id for item in persisted.findings],
        )

    # A second application lifespan proves restart-safe read and export.
    get_settings.cache_clear()
    with TestClient(app) as client:
        assert client.get(
            f"/api/cases/{run.case_id}/runs/{run.run_id}"
        ).json()["finding_count"] == 25
        exported = client.get(f"/api/cases/{run.case_id}/exports/xlsx")
        assert exported.status_code == 200
        xlsx_findings, xlsx_rules, finding_ids = _xlsx_counts(exported.content)
        assert (xlsx_findings, xlsx_rules) == (25, 29)
        assert len(finding_ids) == len(set(finding_ids)) == 25


@pytest.mark.parametrize("empty_value", ["", "   ", "\t"])
def test_finish_running_run_normalizes_optional_parameters(empty_value, tmp_path) -> None:
    run = _review_run()
    run.rule_results[0] = run.rule_results[0].model_copy(update={"parameter": empty_value})
    run.findings[0] = run.findings[0].model_copy(update={"parameter": empty_value})
    runtime = DatabaseRuntime(tmp_path / "review.db")
    runtime.initialize()
    token = _create_running_case(runtime, run)
    session = runtime.session()
    repository = ReviewRepository(session)
    repository.finish_running_run(run, token)
    persisted = repository.get_run(run.run_id)
    assert persisted.rule_results[0].parameter is None
    assert persisted.findings[0].parameter is None
    repository.finish_running_run(run, token)
    assert _database_counts(runtime, run.run_id) == (25, 29)
    session.close()
    runtime.dispose()


def test_finish_running_run_rolls_back_production_equivalent_payload_atomically(tmp_path) -> None:
    run = _review_run()
    runtime = DatabaseRuntime(tmp_path / "review.db")
    runtime.initialize()
    token = _create_running_case(runtime, run)
    session = runtime.session()
    repository = ReviewRepository(session)
    checkpoint = [Finding.model_validate(item) for item in _fixture()["checkpoint_findings"]]
    repository.checkpoint_running_run(run, token, checkpoint[:8])
    run.findings[1] = run.findings[0].model_copy(deep=True)
    with pytest.raises(IntegrityError):
        repository.finish_running_run(run, token)
    session.close()

    restarted = DatabaseRuntime(tmp_path / "review.db")
    restarted.initialize()
    persisted = ReviewRepository(restarted.session()).get_run(run.run_id)
    assert persisted.final_status == "RUNNING"
    assert len(persisted.findings) == 8
    assert len(persisted.rule_results) == 0
    assert _database_counts(restarted, run.run_id) == (8, 0)
    restarted.dispose()
