from pathlib import Path

from docx import Document
import pytest

from app.domain.enums import BlockType
from app.domain.exceptions import ParseError
from app.parsers.docx_parser import DocxParser


def make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("一、项目概况", level=1)
    doc.add_paragraph("本项目位于测试区。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "参数名称"
    table.cell(0, 1).text = "数值"
    table.cell(1, 0).text = "开发井总数"
    table.cell(1, 1).text = "36"
    doc.save(path)


def test_parser_emits_traceable_spans(tmp_path: Path) -> None:
    path = tmp_path / "minimal.docx"
    make_docx(path)

    parsed = DocxParser().parse(path, document_id="D1")

    assert parsed.document_id == "D1"
    assert any(span.block_type is BlockType.HEADING for span in parsed.spans)
    cell = next(span for span in parsed.table_cells if span.text == "36")
    assert (cell.table_index, cell.row_index, cell.column_index) == (0, 1, 1)
    assert cell.section_path == ["一、项目概况"]
    assert cell.text_hash


def test_parser_is_deterministic(tmp_path: Path) -> None:
    path = tmp_path / "minimal.docx"
    make_docx(path)

    first = DocxParser().parse(path, document_id="D1")
    second = DocxParser().parse(path, document_id="D1")

    assert [span.model_dump() for span in first.spans] == [
        span.model_dump() for span in second.spans
    ]


@pytest.mark.parametrize(
    "path",
    [
        pytest.param(lambda tmp_path: tmp_path / "missing.docx", id="missing"),
        pytest.param(
            lambda tmp_path: _write_corrupt_docx(tmp_path / "corrupt.docx"),
            id="corrupt",
        ),
    ],
)
def test_parser_normalizes_unreadable_docx_errors(tmp_path: Path, path) -> None:
    docx_path = path(tmp_path)

    with pytest.raises(ParseError, match="Unable to parse DOCX document") as error:
        DocxParser().parse(docx_path)

    assert error.value.__cause__ is not None
    assert str(docx_path) not in str(error.value)


def _write_corrupt_docx(path: Path) -> Path:
    path.write_bytes(b"this is not a ZIP archive")
    return path


def test_parser_tracks_nested_heading_levels(tmp_path: Path) -> None:
    path = tmp_path / "nested-headings.docx"
    doc = Document()
    doc.add_heading("Top", level=1)
    doc.add_heading("Nested", level=2)
    doc.add_paragraph("Nested body")
    doc.add_heading("Next top", level=1)
    doc.add_paragraph("Top body")
    doc.save(path)

    parsed = DocxParser().parse(path, document_id="D3")

    assert [span.section_path for span in parsed.paragraphs] == [
        ["Top"],
        ["Top", "Nested"],
        ["Top", "Nested"],
        ["Next top"],
        ["Next top"],
    ]


def test_parser_interleaves_tables_with_active_xml_section_path(tmp_path: Path) -> None:
    path = tmp_path / "interleaved.docx"
    doc = Document()
    doc.add_heading("First section", level=1)
    doc.add_paragraph("First body")
    first_table = doc.add_table(rows=1, cols=1)
    first_table.cell(0, 0).text = "first value"
    doc.add_heading("Second section", level=1)
    second_table = doc.add_table(rows=1, cols=1)
    second_table.cell(0, 0).text = "second value"
    doc.add_paragraph("Second body")
    doc.save(path)

    parsed = DocxParser().parse(path, document_id="D2")

    assert [span.span_id for span in parsed.spans] == [
        "D2:p:0",
        "D2:p:1",
        "D2:t:0:0:0",
        "D2:p:2",
        "D2:t:1:0:0",
        "D2:p:3",
    ]
    assert [span.section_path for span in parsed.table_cells] == [
        ["First section"],
        ["Second section"],
    ]
    assert [(span.table_index, span.row_index, span.column_index) for span in parsed.table_cells] == [
        (0, 0, 0),
        (1, 0, 0),
    ]


def test_parser_deduplicates_merged_xml_cells_but_keeps_distinct_equal_text(tmp_path: Path):
    path = tmp_path / "merged.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "merged value"
    table.cell(1, 0).text = "same value"
    table.cell(1, 1).text = "same value"
    doc.save(path)

    parsed = DocxParser().parse(path, document_id="D4")
    assert [span.text for span in parsed.table_cells] == [
        "merged value",
        "same value",
        "same value",
    ]
    assert [span.span_id for span in parsed.table_cells] == [
        "D4:t:0:0:0",
        "D4:t:0:1:0",
        "D4:t:0:1:1",
    ]
