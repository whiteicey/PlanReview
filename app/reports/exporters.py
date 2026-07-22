"""Generate safe, local finding exports without source documents or provider metadata."""

from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from openpyxl import Workbook

from app.domain.enums import FindingCategory, Origin, ReviewStatus, Severity
from app.review.pipeline import ReviewRun
from app.settings import get_settings

_FINDING_COLUMNS = (
    "finding_id", "origin", "category", "severity", "title", "description",
    "suggestion", "location", "evidence_span_ids", "review_status", "human_note",
)
_RULE_COLUMNS = ("rule_id", "category", "message", "evidence_span_ids")
_EVIDENCE_COLUMNS = ("span_id", "evidence_text", "location", "file_name")
_ILLEGAL_EXCEL_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
_HASH_RE = re.compile(r"[0-9a-f]{64}\Z")
_SAFE_LABEL = re.compile(r"[A-Za-z0-9_.:-]{1,128}\Z")
_ALLOWED_CATEGORIES = tuple(FindingCategory)


def _finding_locations(run: ReviewRun, item) -> str:
    """Readable source locations for a finding's evidence spans."""
    seen: list[str] = []
    for span_id in item.evidence_span_ids:
        location = run.evidence_locations.get(span_id)
        if location and location not in seen:
            seen.append(location)
    return "；".join(seen)


def _rows(run: ReviewRun) -> list[dict[str, str | None]]:
    """Represent current finding state, including an expert's persisted review."""
    return [
        {
            "finding_id": item.finding_id,
            "origin": item.origin.value,
            "category": item.category,
            "severity": item.severity.value,
            "title": item.title,
            "description": item.description,
            "suggestion": item.suggestion,
            "location": _finding_locations(run, item),
            "evidence_span_ids": ", ".join(item.evidence_span_ids),
            "review_status": item.review_status.value,
            "human_note": item.human_note,
        }
        for item in run.findings
    ]


def safe_excel_cell(value):
    """Return an Excel-safe external value without changing typed scalars."""
    if not isinstance(value, str):
        return value
    cleaned = _ILLEGAL_EXCEL_CONTROL_CHARS.sub("", value)
    if cleaned.startswith(("=", "+", "-", "@")):
        return "'" + cleaned
    return cleaned


def _append_external_row(sheet, values) -> None:
    sheet.append([safe_excel_cell(value) for value in values])


def _assert_formula_free(workbook: Workbook) -> None:
    """This exporter intentionally has no system-authored formulas."""
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if cell.data_type == "f":
                    raise ValueError("export workbook must not contain formulas")


def export_excel(
    run: ReviewRun,
    target: Path,
    *,
    evidence_texts: dict[str, str] | None = None,
    evidence_file_names: dict[str, str] | None = None,
) -> Path:
    """Write editable review-state rows and evidence references to a spreadsheet."""
    target = _prepare_target(target)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Findings"
    _append_external_row(sheet, [get_settings().disclaimer])
    _append_external_row(sheet, _FINDING_COLUMNS)
    for row in _rows(run):
        _append_external_row(sheet, [row[column] for column in _FINDING_COLUMNS])
    sheet.freeze_panes = "A3"

    rules_sheet = workbook.create_sheet("Rules")
    _append_external_row(rules_sheet, _RULE_COLUMNS)
    for result in run.rule_results:
        _append_external_row(
            rules_sheet,
            (
                result.rule_id,
                result.category,
                result.message,
                ", ".join(result.evidence_span_ids),
            ),
        )

    evidence_sheet = workbook.create_sheet("Evidence")
    _append_external_row(evidence_sheet, _EVIDENCE_COLUMNS)
    texts = evidence_texts or {}
    file_names = evidence_file_names or {}
    referenced_span_ids = dict.fromkeys(
        span_id for finding in run.findings for span_id in finding.evidence_span_ids
    )
    for span_id in referenced_span_ids:
        _append_external_row(
            evidence_sheet,
            (
                span_id,
                texts.get(span_id, ""),
                run.evidence_locations.get(span_id, ""),
                file_names.get(span_id, ""),
            ),
        )

    _assert_formula_free(workbook)
    workbook.save(target)
    return target


def export_word(run: ReviewRun, target: Path) -> Path:
    """Write a human-readable report including evidence and expert review state."""
    target = _prepare_target(target)
    document = Document()
    document.add_heading("审查发现", level=0)
    document.add_paragraph(get_settings().disclaimer)
    document.add_paragraph(f"审查状态：{run.final_status}")
    for item in run.findings:
        document.add_heading(item.title, level=1)
        document.add_paragraph(f"问题编号：{item.finding_id}")
        document.add_paragraph(
            f"来源：{item.origin.value}；严重性：{item.severity.value}；"
            f"专家状态：{item.review_status.value}"
        )
        document.add_paragraph(item.description)
        document.add_paragraph(f"建议：{item.suggestion}")
        location = _finding_locations(run, item)
        document.add_paragraph(f"问题位置：{location or '未定位'}")
        document.add_paragraph(f"证据 span：{', '.join(item.evidence_span_ids) or '无'}")
        if item.human_note:
            document.add_paragraph(f"专家备注：{item.human_note}")
    document.save(target)
    return target


def export_anonymous_package(run: ReviewRun, target_zip: Path) -> Path:
    """Write a strict allow-list ZIP with de-identified findings and no raw sources."""
    target_zip = _prepare_target(target_zip)
    span_aliases = _span_aliases(run)
    hashes = _anonymous_evidence_hashes(run, span_aliases)
    review_counts = Counter(item.review_status.value for item in run.findings)
    payload = {
        "disclaimer": get_settings().disclaimer,
        "findings": [
            {
                "finding_id": f"finding-{index:04d}",
                "origin": _opaque_enum(item.origin, Origin, "origin"),
                "category": _opaque_category(item.category),
                "severity": _opaque_enum(item.severity, Severity, "severity"),
                "evidence_span_ids": [span_aliases[span_id] for span_id in item.evidence_span_ids],
                "review_status": _opaque_enum(item.review_status, ReviewStatus, "review status"),
            }
            for index, item in enumerate(run.findings, start=1)
        ],
        "rule_versions": [
            {
                "rule_id": f"rule-{index:04d}",
                "version": f"version-{index:04d}",
            }
            for index, result in enumerate(_versioned_rule_results(run), start=1)
        ],
        "evidence_text_hashes": hashes,
        "metrics": {
            "finding_count": len(run.findings),
            "review_state_counts": {
                _opaque_enum(ReviewStatus(state), ReviewStatus, "review status"): count
                for state, count in sorted(review_counts.items())
            },
            "accuracy": "not_measured",
            "recall": "not_measured",
            "time_saved": "not_measured",
            "cost": "not_measured",
        },
    }
    with ZipFile(target_zip, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "anonymous-findings.json",
            json.dumps(payload, ensure_ascii=False, separators=(",", ":")),
        )
    return target_zip


def _prepare_target(target: Path) -> Path:
    path = Path(target)
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _span_aliases(run: ReviewRun) -> dict[str, str]:
    aliases: dict[str, str] = {}
    for finding in run.findings:
        for span_id in finding.evidence_span_ids:
            aliases.setdefault(span_id, f"evidence-{len(aliases) + 1:04d}")
    return aliases


def _versioned_rule_results(run: ReviewRun):
    """Yield one version record per rule, preserving first evaluation order."""
    seen: set[str] = set()
    for result in run.rule_results:
        if result.rule_id not in seen and _rule_version(result) is not None:
            seen.add(result.rule_id)
            yield result


def _rule_version(result) -> str | None:
    if result.rule_version is not None:
        return result.rule_version
    version = result.details.get("rule_version") if isinstance(result.details, dict) else None
    return version if isinstance(version, str) else None


def _anonymous_evidence_hashes(run: ReviewRun, aliases: dict[str, str]) -> dict[str, str]:
    """Export only validated source-text digests, never source text or span IDs."""
    hashes: dict[str, str] = {}
    for span_id, alias in aliases.items():
        text_hash = run.evidence_text_hashes.get(span_id)
        if text_hash is None:
            # Legacy or manually constructed runs may lack source hashes. Do not invent one.
            continue
        if not isinstance(text_hash, str) or _HASH_RE.fullmatch(text_hash) is None:
            raise ValueError("evidence text hash must be a lowercase SHA-256 digest")
        hashes[alias] = text_hash
    return hashes


def _opaque_enum(value, enum_type, field_name: str) -> str:
    """Expose only a finite enum member, mapped to an opaque stable alias."""
    try:
        member = value if isinstance(value, enum_type) else enum_type(value)
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{field_name} is not an approved taxonomy value") from exc
    return f"{enum_type.__name__.lower()}-{list(enum_type).index(member) + 1:04d}"


def _opaque_category(value: str | FindingCategory) -> str:
    """Map the finite supported category taxonomy to opaque aliases."""
    try:
        category = value if isinstance(value, FindingCategory) else FindingCategory(value)
    except (TypeError, ValueError) as exc:
        raise ValueError("category is not an approved taxonomy value") from exc
    return f"category-{_ALLOWED_CATEGORIES.index(category) + 1:04d}"
